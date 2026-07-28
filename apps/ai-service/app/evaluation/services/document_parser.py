import io
import logging
from pathlib import Path

import pdfplumber
from docx import Document

from app.evaluation.exceptions import DocumentExtractionError, UnsupportedFileTypeError

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
SUPPORTED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    # Browsers / proxies often send these instead of the precise type
    "application/octet-stream",
    "binary/octet-stream",
}


def _normalize_content_type(content_type: str | None) -> str | None:
    if not content_type:
        return None
    # Strip charset / boundary parameters: "application/pdf; charset=binary"
    return content_type.split(";", 1)[0].strip().lower() or None


class DocumentParser:
    def validate_file(self, filename: str | None, content_type: str | None) -> str:
        if not filename:
            raise UnsupportedFileTypeError("Evaluation file is required.")

        extension = Path(filename).suffix.lower()
        if extension == ".doc":
            raise UnsupportedFileTypeError(
                "Legacy .doc files are not supported. Please upload PDF or DOCX."
            )
        if extension in {".jpg", ".jpeg", ".png", ".webp", ".gif"}:
            raise UnsupportedFileTypeError(
                "Image files are not supported for AI evaluation extraction. "
                "Please upload a text-based PDF or DOCX."
            )
        if extension not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError()

        normalized = _normalize_content_type(content_type)
        if (
            normalized
            and normalized not in SUPPORTED_CONTENT_TYPES
            and not normalized.startswith("application/pdf")
        ):
            # Trust the file extension when MIME is odd but extension is valid
            logger.warning(
                "Unexpected evaluation MIME type %s for %s — continuing by extension",
                content_type,
                filename,
            )

        return extension

    def extract_text(self, file_bytes: bytes, extension: str) -> str:
        if not file_bytes:
            raise DocumentExtractionError("Uploaded evaluation file is empty.")

        try:
            if extension == ".pdf":
                return self._extract_from_pdf(file_bytes)
            return self._extract_from_docx(file_bytes)
        except (DocumentExtractionError, UnsupportedFileTypeError):
            # Preserve specific messages (e.g. "No readable text found…")
            raise
        except Exception as exc:
            raise DocumentExtractionError(
                f"Failed to extract text from evaluation document: {exc}"
            ) from exc

    def _extract_from_pdf(self, file_bytes: bytes) -> str:
        document_text = self._extract_pdf_with_pdfplumber(file_bytes)
        if not document_text:
            document_text = self._extract_pdf_with_pypdf(file_bytes)

        if not document_text:
            raise DocumentExtractionError(
                "No readable text found in PDF. If this is a scanned/image PDF, "
                "export a text-based PDF or DOCX and try again."
            )
        return document_text

    def _extract_pdf_with_pdfplumber(self, file_bytes: bytes) -> str:
        try:
            pages: list[str] = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    # Also pull table cell text when layout extraction misses body copy
                    if not text.strip():
                        tables = page.extract_tables() or []
                        cells: list[str] = []
                        for table in tables:
                            for row in table:
                                for cell in row or []:
                                    if cell and str(cell).strip():
                                        cells.append(str(cell).strip())
                        text = " ".join(cells)
                    if text.strip():
                        pages.append(text.strip())
            return "\n\n".join(pages).strip()
        except Exception as exc:
            logger.warning("pdfplumber failed for evaluation PDF: %s", exc)
            return ""

    def _extract_pdf_with_pypdf(self, file_bytes: bytes) -> str:
        try:
            from pypdf import PdfReader  # optional fallback dependency
        except ImportError:
            logger.warning("pypdf not installed — skipping PDF fallback extractor")
            return ""

        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            if getattr(reader, "is_encrypted", False):
                try:
                    reader.decrypt("")
                except Exception:
                    raise DocumentExtractionError(
                        "PDF is password-protected. Remove the password and try again."
                    ) from None

            pages: list[str] = []
            for page in reader.pages:
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(text)
            return "\n\n".join(pages).strip()
        except DocumentExtractionError:
            raise
        except Exception as exc:
            logger.warning("pypdf fallback failed for evaluation PDF: %s", exc)
            return ""

    def _extract_from_docx(self, file_bytes: bytes) -> str:
        document = Document(io.BytesIO(file_bytes))
        parts: list[str] = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        # Tables often hold scorecards in evaluation templates
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        document_text = "\n".join(parts).strip()
        if not document_text:
            raise DocumentExtractionError(
                "No readable text found in DOCX. Ensure the file contains selectable text."
            )
        return document_text
